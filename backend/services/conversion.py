from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple, Any, Dict

from docx import Document
from langchain_core.documents.base import Document as LangchainDocument
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    CharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    TokenTextSplitter,
    NLTKTextSplitter,
    SpacyTextSplitter,
)

from ..config import settings
from ..schemas import ChunkingMethod

from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import (
    PdfPipelineOptions,
)
from docling.datamodel.accelerator_options import AcceleratorDevice, AcceleratorOptions
from docling.document_converter import DocumentConverter, PdfFormatOption


@dataclass
class ChunkPayload:
    text: str
    section_heading: Optional[str]
    page_number: Optional[int]


def _token_chunks(text: str, heading: Optional[str], page_number: Optional[int]) -> List[ChunkPayload]:
    tokens = text.split()
    size = settings.chunk_size
    overlap = settings.chunk_overlap
    chunks: List[ChunkPayload] = []
    start = 0
    while start < len(tokens):
        end = min(start + size, len(tokens))
        chunk_tokens = tokens[start:end]
        chunk_text = " ".join(chunk_tokens)
        chunks.append(ChunkPayload(text=chunk_text, section_heading=heading, page_number=page_number))
        if end == len(tokens):
            break
        start = max(end - overlap, 0)
    return chunks


def convert_to_chunks(
    path: Path,
    filetype: str,
    chunking_method: ChunkingMethod = ChunkingMethod.RECURSIVE_CHARACTER
) -> Tuple[List[ChunkPayload], bool, str]:
    """Convert file to chunks using specified chunking method. Returns (chunks, used_docling, raw_markdown)."""
    # Convert all file types to markdown first
    if filetype in {"pdf", "docx"}:
        # Use Docling to convert to markdown
        markdown = _convert_with_docling_to_markdown(path)
    else:  # txt
        # Plain text is valid markdown
        markdown = path.read_text(encoding="utf-8", errors="ignore")
    
    # Convert markdown to chunks
    chunks = markdown_to_chunks(markdown, chunking_method)
    return chunks, True, markdown


def _convert_with_docling_to_markdown(path: Path) -> str:
    """Convert a document to Markdown using Docling. Returns empty string on failure."""
    try:
        # Import inside function to avoid import-time errors if Docling isn't installed yet.

        pipeline_options = PdfPipelineOptions()
        accelerator = AcceleratorOptions(device=AcceleratorDevice.CUDA, num_threads=6,)

        pipeline_options.do_ocr = True
        pipeline_options.do_table_structure = True
        pipeline_options.do_code_enrichment = True
        pipeline_options.do_formula_enrichment = True
        pipeline_options.do_picture_description = True
        pipeline_options.accelerator_options = accelerator

        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options),
            }
        )
        # Docling can take a path as source; it auto-detects format.
        result = converter.convert(str(path))
        return result.document.export_to_markdown()
        
    except Exception:
        # Any error (missing package, conversion issue) -> fallback
        return ""


def markdown_to_chunks(
    markdown_text: str,
    chunking_method: ChunkingMethod = ChunkingMethod.RECURSIVE_CHARACTER
) -> List[ChunkPayload]:
    """Chunk Markdown text using the specified LangChain text splitter."""
    if not markdown_text:
        return []
    
    # Create the appropriate text splitter based on method
    docs: List[LangchainDocument] = []
    
    try:
        if chunking_method == ChunkingMethod.MARKDOWN_HEADER:
            # Configure Markdown header-based splitter
            header_config: List[Tuple[str, str]] = [
                ("#", "Header 1"),
                ("##", "Header 2"),
                ("###", "Header 3"),
            ]
            md_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=header_config)
            docs = md_splitter.split_text(markdown_text)
        
        elif chunking_method == ChunkingMethod.CHARACTER:
            char_splitter = CharacterTextSplitter(
                chunk_size=settings.chunk_size,
                chunk_overlap=settings.chunk_overlap,
                separator="\n\n",
            )
            text_chunks = char_splitter.split_text(markdown_text)
            docs = [LangchainDocument(page_content=chunk) for chunk in text_chunks]
        
        elif chunking_method == ChunkingMethod.TOKEN:
            token_splitter = TokenTextSplitter(
                chunk_size=settings.chunk_size,
                chunk_overlap=settings.chunk_overlap,
            )
            text_chunks = token_splitter.split_text(markdown_text)
            docs = [LangchainDocument(page_content=chunk) for chunk in text_chunks]
        
        elif chunking_method == ChunkingMethod.NLTK:
            nltk_splitter = NLTKTextSplitter(
                chunk_size=settings.chunk_size,
                chunk_overlap=settings.chunk_overlap,
            )
            text_chunks = nltk_splitter.split_text(markdown_text)
            docs = [LangchainDocument(page_content=chunk) for chunk in text_chunks]
        
        elif chunking_method == ChunkingMethod.SPACY:
            spacy_splitter = SpacyTextSplitter(
                chunk_size=settings.chunk_size,
                chunk_overlap=settings.chunk_overlap,
            )
            text_chunks = spacy_splitter.split_text(markdown_text)
            docs = [LangchainDocument(page_content=chunk) for chunk in text_chunks]
        
        else:  # RECURSIVE_CHARACTER (default)
            char_splitter = RecursiveCharacterTextSplitter(
                chunk_size=settings.chunk_size,
                chunk_overlap=settings.chunk_overlap,
                is_separator_regex=False,
            )
            text_chunks = char_splitter.split_text(markdown_text)
            docs = [LangchainDocument(page_content=chunk) for chunk in text_chunks]
    
    except Exception as e:
        # Fallback to recursive character splitter if the selected method fails
        print(f"Error with {chunking_method}, falling back to recursive character: {e}")
        char_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            is_separator_regex=False,
        )
        text_chunks = char_splitter.split_text(markdown_text)
        docs = [LangchainDocument(page_content=chunk) for chunk in text_chunks]
    
    # If still no docs, return empty
    if not docs:
        return []

    # Convert Documents to ChunkPayload, using deepest available header as section heading
    chunks: List[ChunkPayload] = []
    header_names_in_order = ["Header 1", "Header 2", "Header 3"]
    for doc in docs:
        content = doc.page_content.strip()
        if not content:
            continue
        heading: Optional[str] = None
        # Prefer deepest header present (e.g., Header 4 over Header 1)
        meta: Dict[str, Any] = getattr(doc, "metadata", {})
        for name in reversed(header_names_in_order):
            val = meta.get(name)
            if val is not None:
                heading = val if isinstance(val, str) else str(val)
                break
        chunks.append(ChunkPayload(text=content, section_heading=heading, page_number=None))

    return chunks
